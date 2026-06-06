"""
TPM Command Center — API Server
================================
Single Flask server that:
  - Serves the dashboard HTML
  - Exposes REST endpoints for all agents
  - Manages jobs.db (SQLite)
  - Streams agent output back to the browser via SSE

Run:  python server.py
Open: http://localhost:5050

Requires: pip install flask flask-cors crewai crewai-tools
          requests beautifulsoup4 python-docx ollama schedule playwright
          playwright install chromium
"""

import json
import os
import queue
import re
import sqlite3
import subprocess
import sys
import threading
import time
from datetime import datetime
from pathlib import Path

from flask import Flask, Response, jsonify, request, send_from_directory
from flask_cors import CORS

# ── Models ───────────────────────────────────────────────────────────────────
# Pull these with Ollama:
#   ollama pull llama3.1:8b           (fast general — chat, agents, discovery)
#   ollama pull deepseek-coder-v2:16b (code edits, structured output, admin panel)
#   ollama pull mistral:7b-instruct   (fast, great at following instructions — resume bullets)
#   ollama pull qwen2.5:14b           (strong writing quality at mid-size — resume/cover letters)
#   ollama pull llama3.1:70b-instruct-q4_K_M  (highest quality — final resume polish, optional)
#   ollama pull deepseek-r1:14b       (reasoning model — great for job analysis + matching)
#
# Recommended for your resume pipeline:
#   Step 1 (extract requirements): mistral:7b-instruct or deepseek-r1:14b
#   Step 2 (write summary):        qwen2.5:14b
#   Step 3 (rewrite bullets):      qwen2.5:14b or deepseek-coder-v2:16b
#   Chat/agents:                    llama3.1:8b (speed) or qwen2.5:14b (quality)
#   Code edits:                     deepseek-coder-v2:16b

MODEL_FAST   = "llama3.1:8b"                   # ~98 tok/s — chat, agents, discovery
MODEL_MID    = "qwen2.5:14b"                   # ~30 tok/s — best quality/speed for writing
MODEL_HQ     = "llama3.1:70b-instruct-q4_K_M"  # ~3 tok/s — final polish only
MODEL_CODE   = "deepseek-coder-v2:16b"          # structured output, code edits
MODEL_REASON = "deepseek-r1:14b"                # reasoning/analysis tasks

# ── Editable files (admin code workspace) ────────────────────────────────────
EDITABLE_FILES = [
    "scraper.py", "server.py", "dashboard.html",
    "ma_scraper.py", "resume_agent.py", "scout_agent.py",
    "scraper_agent.py", "tracker.py", "backfill.py",
    "profile.py", "diagnose.py", "test_google.py",
]

# ── Paths ─────────────────────────────────────────────────────────────────────
BASE_DIR      = Path(__file__).parent
DB_PATH       = BASE_DIR / "jobs.db"
RESUME_DIR    = BASE_DIR / "resumes"
COVER_DIR     = BASE_DIR / "cover_letters"
TAILORED_DIR  = BASE_DIR / "tailored_resumes"
DIGEST_DIR    = BASE_DIR / "digests"

for d in [RESUME_DIR, COVER_DIR, TAILORED_DIR, DIGEST_DIR]:
    d.mkdir(exist_ok=True)

app = Flask(__name__, static_folder=str(BASE_DIR), static_url_path="")
CORS(app)

# SSE queues per task_id
sse_queues: dict[str, queue.Queue] = {}


# ── Database ──────────────────────────────────────────────────────────────────

def get_db():
    con = sqlite3.connect(str(DB_PATH))
    con.row_factory = sqlite3.Row
    return con


def init_db():
    con = get_db()
    con.executescript("""
        CREATE TABLE IF NOT EXISTS jobs (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            company     TEXT NOT NULL,
            title       TEXT NOT NULL,
            location    TEXT,
            work_type   TEXT,
            level       TEXT,
            team        TEXT,
            salary_min  INTEGER,
            salary_max  INTEGER,
            salary_raw  TEXT,
            posted_date TEXT,
            description TEXT,
            url         TEXT UNIQUE,
            source      TEXT,
            job_id      TEXT,
            found_date  TEXT,
            status      TEXT DEFAULT 'new',
            score       INTEGER,
            notes       TEXT
        );
        CREATE TABLE IF NOT EXISTS resumes (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            name        TEXT NOT NULL,
            company     TEXT,
            role        TEXT,
            filename    TEXT NOT NULL,
            type        TEXT DEFAULT 'tailored',
            created     TEXT,
            job_id      INTEGER
        );
        CREATE TABLE IF NOT EXISTS cover_letters (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            name        TEXT NOT NULL,
            company     TEXT,
            role        TEXT,
            filename    TEXT NOT NULL,
            created     TEXT,
            job_id      INTEGER
        );
        CREATE TABLE IF NOT EXISTS interview_prep (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            company     TEXT NOT NULL,
            role        TEXT,
            content     TEXT,
            created     TEXT,
            job_id      INTEGER
        );
    """)
    # Migrate existing DB — safely add new columns if missing
    for col, typedef in [
        ("work_type",   "TEXT"),
        ("level",       "TEXT"),
        ("team",        "TEXT"),
        ("salary_min",  "INTEGER"),
        ("salary_max",  "INTEGER"),
        ("salary_raw",  "TEXT"),
        ("posted_date", "TEXT"),
        ("description", "TEXT"),
        ("score",       "INTEGER"),
    ]:
        try:
            con.execute(f"ALTER TABLE jobs ADD COLUMN {col} {typedef}")
        except Exception:
            pass
    con.commit()
    con.close()


def rows_to_list(rows) -> list:
    return [dict(r) for r in rows] if rows else []


# ── SSE helpers ───────────────────────────────────────────────────────────────

def sse_send(task_id: str, event: str, data: dict):
    # Auto-create queue if missing (handles race condition where thread
    # starts before browser connects to /api/stream/<task_id>)
    if task_id not in sse_queues:
        sse_queues[task_id] = queue.Queue()
    sse_queues[task_id].put({"event": event, "data": data})


def sse_stream(task_id: str):
    # Register queue — if thread already ran, drain existing messages
    if task_id not in sse_queues:
        sse_queues[task_id] = queue.Queue()
    q = sse_queues[task_id]
    try:
        while True:
            try:
                msg = q.get(timeout=60)
                yield f"event: {msg['event']}\ndata: {json.dumps(msg['data'])}\n\n"
                if msg["event"] in ("done", "error"):
                    break
            except queue.Empty:
                yield "event: ping\ndata: {}\n\n"
    finally:
        sse_queues.pop(task_id, None)


# ── Routes: Dashboard ─────────────────────────────────────────────────────────

@app.route("/")
def index():
    return send_from_directory(str(BASE_DIR), "dashboard.html")


# ── Routes: Recruiters ────────────────────────────────────────────────────────

@app.route("/api/recruiters", methods=["GET"])
def get_recruiters():
    company = request.args.get("company", "")
    con = get_db()
    try:
        con.execute("""
            CREATE TABLE IF NOT EXISTS recruiters (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                company TEXT, domain TEXT, name TEXT,
                first_name TEXT, last_name TEXT, email TEXT,
                email_status TEXT, title TEXT, linkedin_url TEXT,
                source TEXT, confidence INTEGER, found_date TEXT,
                UNIQUE(company, email)
            )
        """)
        sql = "SELECT * FROM recruiters WHERE 1=1"
        params = []
        if company:
            sql += " AND company LIKE ?"; params.append(f"%{company}%")
        sql += " ORDER BY confidence DESC, found_date DESC"
        rows = rows_to_list(con.execute(sql, params).fetchall())
        con.close()
        return jsonify(rows)
    except Exception as e:
        con.close()
        return jsonify({"error": str(e)}), 500


@app.route("/api/recruiters/run", methods=["POST"])
def trigger_recruiter_lookup():
    """Trigger recruiter enrichment for a specific company or top unscraped."""
    d = request.json or {}
    company = d.get("company")
    limit   = d.get("limit", 10)
    try:
        from recruiter_lookup import run as run_recruiters
        saved = run_recruiters(company_filter=company, limit=limit, verbose=False)
        return jsonify({"status": "ok", "saved": saved})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ── Routes: Jobs ──────────────────────────────────────────────────────────────

@app.route("/api/jobs", methods=["GET"])
def get_jobs():
    status = request.args.get("status")
    search = request.args.get("q", "")
    sql = "SELECT * FROM jobs WHERE 1=1"
    params = []
    if status and status != "all":
        sql += " AND status=?"; params.append(status)
    if search:
        sql += " AND (company LIKE ? OR title LIKE ?)"; params += [f"%{search}%", f"%{search}%"]
    # Exclude jobs with a confirmed salary ceiling below $170k (keep nulls)
    sql += " AND (salary_max IS NULL OR salary_max = 0 OR salary_max >= 170000)"
    sql += " ORDER BY found_date DESC, id DESC"
    con = get_db()
    rows = rows_to_list(con.execute(sql, params).fetchall())
    con.close()
    return jsonify(rows)


@app.route("/api/jobs", methods=["POST"])
def add_job():
    d = request.json or {}
    con = get_db()
    try:
        con.execute(
            "INSERT INTO jobs (company,title,location,url,source,found_date,status,notes) VALUES (?,?,?,?,?,?,?,?)",
            (d.get("company",""), d.get("title",""), d.get("location",""),
             d.get("url",""), d.get("source","Manual"),
             datetime.now().strftime("%Y-%m-%d"), d.get("status","new"), d.get("notes",""))
        )
        con.commit()
    except sqlite3.IntegrityError:
        con.close()
        return jsonify({"error": "URL already exists"}), 409
    con.close()
    return jsonify({"ok": True})


@app.route("/api/jobs/<int:job_id>", methods=["PATCH"])
def update_job(job_id):
    d = request.json or {}
    allowed = {"status", "notes", "score", "title", "location"}
    updates = {k: v for k, v in d.items() if k in allowed}
    if updates:
        set_clause = ", ".join(f"{k}=?" for k in updates)
        con = get_db()
        con.execute(f"UPDATE jobs SET {set_clause} WHERE id=?", list(updates.values()) + [job_id])
        con.commit(); con.close()
    return jsonify({"ok": True})


@app.route("/api/jobs/<int:job_id>", methods=["DELETE"])
def delete_job(job_id):
    con = get_db()
    con.execute("DELETE FROM jobs WHERE id=?", (job_id,))
    con.commit(); con.close()
    return jsonify({"ok": True})


@app.route("/api/stats", methods=["GET"])
def get_stats():
    con = get_db()
    def count(sql, params=()):
        return con.execute(sql, params).fetchone()[0]
    total      = count("SELECT COUNT(*) FROM jobs")
    new_today  = count("SELECT COUNT(*) FROM jobs WHERE found_date=?", (datetime.now().strftime("%Y-%m-%d"),))
    applied    = count("SELECT COUNT(*) FROM jobs WHERE status IN ('applied','interview','offer','rejected')")
    interviews = count("SELECT COUNT(*) FROM jobs WHERE status='interview'")
    offers     = count("SELECT COUNT(*) FROM jobs WHERE status='offer'")
    by_status  = rows_to_list(con.execute("SELECT status, COUNT(*) as count FROM jobs GROUP BY status").fetchall())
    by_company = rows_to_list(con.execute("SELECT company, COUNT(*) as count FROM jobs GROUP BY company ORDER BY count DESC LIMIT 8").fetchall())
    con.close()
    return jsonify({"total": total, "new_today": new_today, "applied": applied,
                    "interviews": interviews, "offers": offers,
                    "by_status": by_status, "by_company": by_company})


# ── Routes: Resumes ───────────────────────────────────────────────────────────

@app.route("/api/resumes", methods=["GET"])
def get_resumes():
    con = get_db()
    rows = rows_to_list(con.execute("SELECT * FROM resumes ORDER BY created DESC").fetchall())
    con.close()
    return jsonify(rows)


@app.route("/api/resumes", methods=["POST"])
def upload_resume():
    if "file" not in request.files:
        return jsonify({"error": "no file"}), 400
    f = request.files["file"]
    filename = f.filename
    f.save(str(RESUME_DIR / filename))
    con = get_db()
    con.execute("INSERT INTO resumes (name,company,role,filename,type,created) VALUES (?,?,?,?,?,?)",
                (request.form.get("name", filename), request.form.get("company",""),
                 request.form.get("role",""), filename,
                 request.form.get("type","master"), datetime.now().strftime("%Y-%m-%d")))
    con.commit(); con.close()
    return jsonify({"ok": True})


@app.route("/api/resumes/<int:rid>/download")
def download_resume(rid):
    con = get_db()
    row = con.execute("SELECT * FROM resumes WHERE id=?", (rid,)).fetchone()
    con.close()
    if not row: return jsonify({"error": "not found"}), 404
    return send_from_directory(str(RESUME_DIR), row["filename"], as_attachment=True)


@app.route("/api/resumes/<int:rid>", methods=["DELETE"])
def delete_resume(rid):
    con = get_db()
    row = con.execute("SELECT * FROM resumes WHERE id=?", (rid,)).fetchone()
    if row:
        try: (RESUME_DIR / row["filename"]).unlink()
        except: pass
        con.execute("DELETE FROM resumes WHERE id=?", (rid,))
        con.commit()
    con.close()
    return jsonify({"ok": True})


# ── Resume/Cover Letter HTML Preview ─────────────────────────────────────────

@app.route("/api/preview", methods=["POST"])
def preview_document():
    """
    Render resume or cover letter text as a clean, print-ready HTML page.
    Input: {content: "markdown-ish text", title: "Resume — Google", type: "resume"|"cover"}
    Returns: Full HTML page suitable for printing, copy-paste to Google Docs, or Save As PDF.
    """
    d = request.json or {}
    content = d.get("content", "")
    title   = d.get("title", "Document")
    doc_type = d.get("type", "resume")

    # Convert the markdown-ish format to HTML
    lines = content.split("\n")
    html_parts = []
    for line in lines:
        line = line.strip()
        if not line:
            html_parts.append("<br>")
        elif line.startswith("## "):
            html_parts.append(f'<h2>{_esc(line[3:])}</h2>')
        elif line.startswith("# "):
            html_parts.append(f'<h1>{_esc(line[2:])}</h1>')
        elif line.startswith("- ") or line.startswith("• "):
            html_parts.append(f'<li>{_esc(line[2:])}</li>')
        elif line.startswith("**") and line.endswith("**"):
            html_parts.append(f'<p><strong>{_esc(line[2:-2])}</strong></p>')
        else:
            html_parts.append(f'<p>{_esc(line)}</p>')

    # Wrap consecutive <li> in <ul>
    body = "\n".join(html_parts)
    body = re.sub(r'(<li>.*?</li>\n?)+', lambda m: f'<ul>{m.group(0)}</ul>', body, flags=re.DOTALL)

    page = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>{_esc(title)}</title>
<style>
@import url('https://fonts.googleapis.com/css2?family=Source+Sans+3:wght@300;400;600;700&display=swap');
* {{ box-sizing: border-box; margin: 0; padding: 0; }}
body {{
    font-family: 'Source Sans 3', 'Segoe UI', sans-serif;
    font-size: 11pt;
    line-height: 1.6;
    color: #1a1a1a;
    max-width: 8.5in;
    margin: 0 auto;
    padding: 0.75in 0.85in;
    background: #fff;
}}
h1 {{
    font-size: 18pt;
    font-weight: 700;
    margin-bottom: 4px;
    color: #111;
    text-align: center;
}}
h2 {{
    font-size: 12pt;
    font-weight: 600;
    text-transform: uppercase;
    letter-spacing: 0.04em;
    border-bottom: 1.5px solid #333;
    padding-bottom: 3px;
    margin: 16px 0 8px;
    color: #222;
}}
p {{
    margin: 3px 0;
}}
strong {{
    font-weight: 600;
}}
ul {{
    margin: 4px 0 4px 20px;
    padding: 0;
}}
li {{
    margin: 3px 0;
    padding-left: 2px;
}}
br {{
    display: block;
    content: "";
    margin: 4px 0;
}}
.toolbar {{
    position: fixed;
    top: 0;
    left: 0;
    right: 0;
    background: #1a1a2e;
    color: #fff;
    padding: 10px 20px;
    display: flex;
    align-items: center;
    gap: 12px;
    font-family: 'Source Sans 3', sans-serif;
    font-size: 13px;
    z-index: 100;
    box-shadow: 0 2px 8px rgba(0,0,0,0.3);
}}
.toolbar button {{
    padding: 6px 16px;
    border-radius: 6px;
    border: 1px solid rgba(255,255,255,0.2);
    background: rgba(255,255,255,0.1);
    color: #fff;
    cursor: pointer;
    font-size: 12px;
    font-family: inherit;
    transition: background 0.15s;
}}
.toolbar button:hover {{
    background: rgba(255,255,255,0.2);
}}
.toolbar button.primary {{
    background: #5b9cf6;
    border-color: #5b9cf6;
}}
.toolbar button.primary:hover {{
    background: #4a8ae0;
}}
.toolbar .spacer {{ flex: 1; }}
.toolbar .title {{ font-weight: 600; }}
@media print {{
    .toolbar {{ display: none; }}
    body {{ padding: 0.5in 0.75in; }}
}}
</style>
</head>
<body>
<div class="toolbar">
    <span class="title">{_esc(title)}</span>
    <span class="spacer"></span>
    <button onclick="copyRichText()">Copy rich text</button>
    <button onclick="window.print()" class="primary">Print / Save PDF</button>
    <span style="font-size:11px;opacity:0.6;">Tip: Paste into Google Docs to edit with formatting</span>
</div>
<div style="margin-top:52px;" id="doc-content">
{body}
</div>
<script>
function copyRichText() {{
    const content = document.getElementById('doc-content');
    const range = document.createRange();
    range.selectNodeContents(content);
    const sel = window.getSelection();
    sel.removeAllRanges();
    sel.addRange(range);
    document.execCommand('copy');
    sel.removeAllRanges();
    const btn = event.target;
    btn.textContent = 'Copied!';
    setTimeout(() => btn.textContent = 'Copy rich text', 2000);
}}
</script>
</body>
</html>"""
    return page, 200, {"Content-Type": "text/html"}


def _esc(s):
    """HTML-escape a string."""
    return str(s or "").replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")


# ── Routes: Cover Letters ─────────────────────────────────────────────────────

@app.route("/api/covers", methods=["GET"])
def get_covers():
    con = get_db()
    rows = rows_to_list(con.execute("SELECT * FROM cover_letters ORDER BY created DESC").fetchall())
    con.close()
    return jsonify(rows)


@app.route("/api/covers/<int:cid>/download")
def download_cover(cid):
    con = get_db()
    row = con.execute("SELECT * FROM cover_letters WHERE id=?", (cid,)).fetchone()
    con.close()
    if not row: return jsonify({"error": "not found"}), 404
    return send_from_directory(str(COVER_DIR), row["filename"], as_attachment=True)


@app.route("/api/covers/<int:cid>", methods=["DELETE"])
def delete_cover(cid):
    con = get_db()
    row = con.execute("SELECT * FROM cover_letters WHERE id=?", (cid,)).fetchone()
    if row:
        try: (COVER_DIR / row["filename"]).unlink()
        except: pass
        con.execute("DELETE FROM cover_letters WHERE id=?", (cid,))
        con.commit()
    con.close()
    return jsonify({"ok": True})


# ── Routes: Interview Prep ────────────────────────────────────────────────────

@app.route("/api/interview", methods=["GET"])
def get_interview_prep():
    con = get_db()
    rows = rows_to_list(con.execute("SELECT * FROM interview_prep ORDER BY created DESC").fetchall())
    con.close()
    return jsonify(rows)


# ── Routes: Agent Tasks (SSE streaming) ───────────────────────────────────────

def run_agent_thread(task_id: str, fn, *args):
    """Run agent function in background thread, stream progress via SSE."""
    def worker():
        try:
            fn(task_id, *args)
        except Exception as e:
            sse_send(task_id, "error", {"message": str(e)})
    t = threading.Thread(target=worker, daemon=True)
    t.start()


@app.route("/api/stream/<task_id>")
def stream(task_id):
    return Response(sse_stream(task_id),
                    mimetype="text/event-stream",
                    headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


# ── Agent: Scraper ────────────────────────────────────────────────────────────

@app.route("/api/agent/heal", methods=["POST"])
def agent_heal():
    """Run the parser healer for a specific company or all companies."""
    d = request.json or {}
    company = d.get("company", "")
    task_id = f"heal_{int(time.time())}"
    run_agent_thread(task_id, _run_healer, company)
    return jsonify({"task_id": task_id})


def _run_healer(task_id: str, company: str):
    sse_send(task_id, "progress", {"step": f"Analyzing {company or 'all'} scrapers...", "pct": 10})
    try:
        sys.path.insert(0, str(BASE_DIR))
        import importlib
        import parser_healer as ph
        importlib.reload(ph)

        if company and company in ph.COMPANY_CONFIG:
            sse_send(task_id, "progress", {"step": f"Fetching live {company} page...", "pct": 25})
            success = ph.heal_company(company, max_attempts=3)
            msg = f"{'Fixed' if success else 'Could not fix'} {company} scraper"
        else:
            sse_send(task_id, "progress", {"step": "Checking all company scrapers...", "pct": 20})
            ph.heal_all(max_attempts=2)
            msg = "Healer run complete — check logs"

        sse_send(task_id, "done", {"message": msg})
    except Exception as e:
        sse_send(task_id, "error", {"message": str(e)})


@app.route("/api/agent/scrape", methods=["POST"])
def agent_scrape():
    task_id = f"scrape_{int(time.time())}"
    run_agent_thread(task_id, _run_scraper)
    return jsonify({"task_id": task_id})


def _run_scraper(task_id: str):
    sse_send(task_id, "progress", {"step": "Starting scraper...", "pct": 5})
    try:
        # Import and run the scraper inline
        sys.path.insert(0, str(BASE_DIR))
        from scraper import init_db as scraper_init, run_scrape
        scraper_init()
        sse_send(task_id, "progress", {"step": "Fetching Google, Amazon, Microsoft, Meta, Apple, NetApp, YC...", "pct": 15})

        # Capture new job count by checking DB before/after
        con = get_db()
        before = con.execute("SELECT COUNT(*) FROM jobs").fetchone()[0]
        con.close()

        # run_scrape now includes LinkedIn, MA scraper, and custom targets
        sse_send(task_id, "progress", {"step": "Running all scrapers including LinkedIn...", "pct": 30})
        run_scrape()

        con = get_db()
        after = con.execute("SELECT COUNT(*) FROM jobs").fetchone()[0]
        con.close()

        new_count = after - before
        sse_send(task_id, "done", {"message": f"Scrape complete — {new_count} new jobs found", "new": new_count})
    except Exception as e:
        sse_send(task_id, "error", {"message": f"Scraper error: {e}"})


# ── Agent: Resume Tailor ──────────────────────────────────────────────────────

@app.route("/api/agent/tailor", methods=["POST"])
def agent_tailor():
    d = request.json or {}
    job_url     = d.get("url", "")
    resume_file = d.get("resume", "")
    notes       = d.get("notes", "")
    if not job_url:
        return jsonify({"error": "url required"}), 400
    task_id = f"tailor_{int(time.time())}"
    quality = d.get("quality", "mid")  # fast, mid, hq
    high_quality = quality == "hq" or d.get("high_quality", False)
    run_agent_thread(task_id, _run_tailor, job_url, resume_file, notes, high_quality, quality)
    return jsonify({"task_id": task_id})


def _run_tailor(task_id: str, job_url: str, resume_file: str, notes: str,
                high_quality: bool = False, quality: str = "mid"):
    """
    Multi-step resume tailor pipeline.
    Breaks the task into focused steps that even a small model handles well:
      1. Extract requirements from job posting
      2. Generate a targeted summary
      3. Rewrite experience bullets
      4. Assemble with template (no LLM — deterministic)
    """
    sse_send(task_id, "progress", {"step": "Fetching job posting...", "pct": 5})
    try:
        import requests as req
        from bs4 import BeautifulSoup
        import ollama

        # Model selection per step based on quality tier
        # fast:  llama3.1:8b for everything (speed priority)
        # mid:   deepseek-r1 for analysis, qwen2.5 for writing (best balance)
        # hq:    deepseek-r1 for analysis, llama3.1:70b for writing (best quality)
        if quality == "fast":
            model_extract = MODEL_FAST
            model_write   = MODEL_FAST
        elif quality == "hq":
            model_extract = MODEL_REASON
            model_write   = MODEL_HQ
        else:  # mid (default)
            model_extract = MODEL_REASON
            model_write   = MODEL_MID

        # ── Step 0: Fetch job posting ──
        headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"}
        r = req.get(job_url, headers=headers, timeout=15)
        soup = BeautifulSoup(r.text, "html.parser")
        for tag in soup(["script","style","nav","footer","header"]): tag.decompose()
        job_text = soup.get_text(separator="\n", strip=True)[:3500]

        # Extract company
        company = "Target Company"
        for c in ["google","amazon","microsoft","meta","apple","netapp","nvidia",
                   "bose","hubspot","akamai","wayfair","toast","klaviyo","dell",
                   "raytheon","cisco","ibm","fidelity","draftkings","oracle",
                   "state street","athenahealth","ptc","rapid7","constant contact"]:
            if c in job_url.lower():
                company = c.title()
                break
        if company == "Target Company":
            title_tag = soup.find("title")
            if title_tag:
                for part in title_tag.get_text().split("|") + title_tag.get_text().split("-"):
                    p = part.strip()
                    if p and len(p) < 40 and not any(kw in p.lower() for kw in ["job","career","apply","search"]):
                        company = p
                        break

        # Read master resume if provided
        resume_content = ""
        if resume_file:
            resume_path = RESUME_DIR / resume_file
            if resume_path.exists():
                try:
                    from docx import Document
                    doc = Document(str(resume_path))
                    resume_content = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                except Exception:
                    pass

        # ── Devon's experience inventory ──
        DEVON_BG = """Devon O'Rourke — Senior Technical Program Manager | Bellingham, MA

Role 1: Senior TPM at Telecom/Infrastructure Company
- Drove end-to-end architecture design and delivery for carrier-grade network infrastructure
- Led cross-functional team of 15+ engineers through full lifecycle: requirements, design, lab testing, UAT, staging, pilot, production
- Owned product roadmap for infrastructure platform serving 2M+ subscribers
- Built and ran lab testing environments; managed UAT cycles reducing production defects by 40%
- Delivered executive reporting and board presentations on program status, risks, milestones
- Managed $3M+ program budgets with vendor negotiations and SOW oversight
- Drove incident escalation and debugging across hardware and software stacks
- Implemented Agile/Scrum across 3 delivery teams, improving sprint velocity 25%

Role 2: Program Manager at Technology Company
- Managed cross-functional delivery of platform migration affecting 500K+ users
- Coordinated 5 engineering teams (backend, frontend, infra, QA, security)
- Established CI/CD pipeline standards reducing deployment time from 4 hours to 20 minutes
- Created stakeholder communication frameworks: weekly exec summaries, quarterly business reviews
- Led pilot programs for 3 new product features with structured rollout to beta customers
- Drove compliance and security reviews for SOC2 and FedRAMP readiness

Role 3: Technical Project Manager (Earlier Career)
- Managed hardware/software integration projects for enterprise networking products
- Ran lab environments for system validation and carrier certification testing
- Coordinated with 3rd-party vendors and contract manufacturers

Skills: JIRA, Confluence, Agile/Scrum, CI/CD, Python, SQL, Terraform, AWS,
Network Architecture, Telecom (OSS/BSS), Stakeholder Management, Executive Reporting,
Vendor Management, Budgeting, Lab Testing, UAT, Product Roadmaps, Risk Management"""

        def _llm(prompt, temp=0.2, max_tok=800, use_model=None):
            """Focused LLM call with model fallback."""
            m = use_model or MODEL_FAST
            try:
                resp = ollama.chat(
                    model=m,
                    messages=[{"role": "user", "content": prompt}],
                    options={"temperature": temp, "num_predict": max_tok},
                    stream=False,
                )
                return resp["message"]["content"].strip()
            except Exception as e:
                # If the preferred model isn't available, fall back to MODEL_FAST
                if m != MODEL_FAST:
                    try:
                        resp = ollama.chat(
                            model=MODEL_FAST,
                            messages=[{"role": "user", "content": prompt}],
                            options={"temperature": temp, "num_predict": max_tok},
                            stream=False,
                        )
                        return resp["message"]["content"].strip()
                    except Exception:
                        pass
                raise e

        # ── Step 1: Extract requirements ──
        sse_send(task_id, "progress", {"step": "Step 1/4 — Extracting requirements from posting...", "pct": 15})

        requirements = _llm(f"""Extract the key requirements from this job posting. Be precise — use the exact words from the posting.

JOB POSTING:
{job_text[:2200]}

Return this EXACT format (one item per line):
TITLE: (exact job title from posting)
COMPANY: (company name)
MUST_HAVE: (8-10 required skills, comma-separated, using exact words from posting)
NICE_TO_HAVE: (3-5 preferred qualifications, comma-separated)
TOOLS: (specific tools/platforms/technologies mentioned)
TEAM: (one sentence about the team this role sits in)
VERBS: (action verbs used in the posting like drive, own, lead, architect, scale)

Output ONLY this list.""", temp=0.1, max_tok=400, use_model=model_extract)

        # Parse
        job_title = must_have = nice_to_have = tools = team_ctx = key_verbs = ""
        for line in requirements.split("\n"):
            l = line.strip()
            if ":" not in l:
                continue
            key, val = l.split(":", 1)
            key = key.strip().upper().replace(" ", "_")
            val = val.strip()
            if key == "TITLE": job_title = val
            elif key == "COMPANY" and val: company = val if company == "Target Company" else company
            elif key in ("MUST_HAVE", "MUST"): must_have = val
            elif key in ("NICE_TO_HAVE", "NICE"): nice_to_have = val
            elif key == "TOOLS": tools = val
            elif key in ("TEAM", "TEAM_CONTEXT"): team_ctx = val
            elif key in ("VERBS", "KEY_VERBS"): key_verbs = val

        # ── Step 2: Targeted summary ──
        sse_send(task_id, "progress", {"step": "Step 2/4 — Writing targeted summary...", "pct": 30})

        summary = _llm(f"""Write a 2-sentence professional summary for a resume. This is for Devon O'Rourke applying to: {job_title} at {company}.

The summary must:
- Start with "Senior Technical Program Manager with 10+ years..."
- Include 3-4 keywords from this list: {must_have}
- Reference the most relevant domain experience (telecom/infrastructure/platform)
- End with what Devon delivers: outcomes, not just activities

2 sentences only. No "I". No fluff. Output ONLY the summary.""", temp=0.3, max_tok=150, use_model=model_write)

        # ── Step 3: Rewrite experience bullets ──
        sse_send(task_id, "progress", {"step": "Step 3/4 — Rewriting experience bullets...", "pct": 50})

        experience = _llm(f"""Rewrite these work experience entries for a resume targeting: {job_title} at {company}

KEYWORDS FROM JOB POSTING (mirror these exactly): {must_have}
ACTION VERBS FROM POSTING: {key_verbs or "drove, led, owned, delivered, scaled, architected"}
TOOLS TO MENTION: {tools}

RAW EXPERIENCE TO REWRITE:
{DEVON_BG}

{f"CURRENT RESUME (adapt from this):{chr(10)}{resume_content[:1200]}" if resume_content else ""}

{f"NOTES: {notes}" if notes else ""}

RULES:
- Output 2-3 roles, 4-5 bullets each
- Every bullet: action verb + what was done + metric (number, percentage, dollar amount, or team size)
- Use the EXACT phrases from the job posting where Devon's experience matches
- Format:
**Company Name — Role Title**
dates
- Bullet with metric
- Bullet with metric

Start IMMEDIATELY with **Company — Role**. No headers. No preamble. No explanation.""",
            temp=0.25, max_tok=1200, use_model=model_write)

        # ── Step 4: Assemble (deterministic — no LLM) ──
        sse_send(task_id, "progress", {"step": "Step 4/4 — Assembling final resume...", "pct": 75})

        summary = _strip_llm_noise(summary).strip()
        experience = _strip_llm_noise(experience).strip()

        # Build skills from job requirements + Devon's actual skills
        job_skills = [s.strip() for s in (must_have + "," + tools).split(",") if s.strip()]
        devon_skills = ["JIRA","Confluence","Agile/Scrum","CI/CD","Python","SQL",
                        "AWS","Terraform","Network Architecture","Telecom","OSS/BSS",
                        "Stakeholder Management","Executive Reporting","Vendor Management",
                        "Lab Testing","UAT","Product Roadmaps","Risk Management"]
        seen = set()
        skills = []
        for s in job_skills + devon_skills:
            s = s.strip()
            if s and s.lower() not in seen and len(s) > 1:
                skills.append(s)
                seen.add(s.lower())
        skills_line = ", ".join(skills[:18])

        from profile import NAME, LOCATION, EMAIL
        linkedin_slug = NAME.lower().replace(" ", "")
        tailored_text = f"""# {NAME}
{LOCATION} | {EMAIL} | linkedin.com/in/{linkedin_slug}

## Summary
{summary}

## Experience
{experience}

## Technical Skills
{skills_line}

## Education
B.S. — (University Name)"""

        # ── Keyword scoring ──
        sse_send(task_id, "progress", {"step": "Scoring keyword coverage...", "pct": 90})

        resume_lower = tailored_text.lower()
        kw_list = [k.strip().lower() for k in must_have.split(",") if k.strip()]
        if not kw_list:
            kw_list = ["program manager","cross-functional","stakeholder","roadmap",
                       "infrastructure","architecture","agile","delivery"]
        matched = [k for k in kw_list if k in resume_lower]
        missed = [k for k in kw_list if k not in resume_lower]
        score = min(100, int((len(matched) / max(len(kw_list), 1)) * 100))

        # Save to DB
        safe = re.sub(r"[^\w]", "_", company)
        date_str = datetime.now().strftime("%Y%m%d_%H%M")
        try:
            con = get_db()
            con.execute(
                "INSERT INTO resumes (name,company,role,filename,type,created) VALUES (?,?,?,?,?,?)",
                (f"Resume — {company}", company, job_title, f"tailored_{safe}_{date_str}",
                 "tailored", datetime.now().strftime("%Y-%m-%d"))
            )
            con.commit(); con.close()
        except Exception:
            pass

        sse_send(task_id, "done", {
            "content": tailored_text,
            "keywords": matched,
            "missed_keywords": missed,
            "score": score,
            "company": company,
            "role": job_title,
            "message": f"Resume tailored for {job_title or 'role'} at {company} — {score}% keyword match"
        })

    except Exception as e:
        sse_send(task_id, "error", {"message": str(e)})


def _strip_llm_noise(text: str) -> str:
    """Remove LLM preamble, postamble, and chatbot pleasantries from generated content."""
    lines = text.strip().split("\n")

    # Strip leading noise — anything before the first # header or ** bold line
    start = 0
    for i, line in enumerate(lines):
        stripped = line.strip()
        if stripped.startswith("#") or stripped.startswith("**") or stripped.startswith("- "):
            start = i
            break
        # Also accept lines that look like contact info (email, phone, location)
        if any(kw in stripped.lower() for kw in ["@", "linkedin", "ma |", "ct |", "ny |"]):
            start = i
            break
    lines = lines[start:]

    # Strip trailing noise — chatbot sign-offs
    noise_phrases = [
        "i hope", "let me know", "good luck", "feel free", "happy to help",
        "i've tailored", "this revised", "this resume", "note:", "changes made",
        "key changes", "here is", "here's", "please review", "best regards",
        "adjustments include", "i made the following", "this version",
    ]
    while lines:
        last = lines[-1].strip().lower()
        if not last:
            lines.pop()
            continue
        if any(last.startswith(p) for p in noise_phrases):
            lines.pop()
            continue
        if last.startswith("---") or last.startswith("***"):
            lines.pop()
            continue
        break

    # Strip markdown code fences
    result = "\n".join(lines)
    result = re.sub(r'^```(?:markdown)?\n?', '', result, flags=re.MULTILINE)
    result = re.sub(r'\n?```$', '', result, flags=re.MULTILINE)

    return result.strip()


# ── Agent: Cover Letter ───────────────────────────────────────────────────────

@app.route("/api/agent/cover", methods=["POST"])
def agent_cover():
    d = request.json or {}
    task_id = f"cover_{int(time.time())}"
    run_agent_thread(task_id, _run_cover, d.get("url",""), d.get("company",""), d.get("role",""))
    return jsonify({"task_id": task_id})


def _run_cover(task_id: str, job_url: str, company: str, role: str):
    sse_send(task_id, "progress", {"step": "Fetching job details...", "pct": 15})
    try:
        job_text = ""
        if job_url:
            import requests as req
            from bs4 import BeautifulSoup
            r = req.get(job_url, headers={"User-Agent":"Mozilla/5.0"}, timeout=15)
            soup = BeautifulSoup(r.text,"html.parser")
            for t in soup(["script","style","nav","footer"]): t.decompose()
            job_text = soup.get_text(separator="\n",strip=True)[:2500]

        sse_send(task_id, "progress", {"step": "Writing cover letter...", "pct": 40})

        import ollama
        prompt = f"""You are a cover letter generator. Output ONLY the cover letter text — no commentary before or after.

Write AS Devon O'Rourke, in first person. This is Devon writing to the hiring manager.

DEVON O'ROURKE — Senior Technical Program Manager
Bellingham, MA | Open to remote, hybrid, or on-site in MA
Background: E2E architecture, infrastructure, telecom, product roadmap delivery,
lab testing/UAT/staging/pilot programs, executive reporting, cross-functional leadership

APPLYING TO:
Company: {company}
Role: {role}
{f"Job details: {job_text[:1500]}" if job_text else ""}

COVER LETTER REQUIREMENTS:
- 3-4 paragraphs, professional but confident tone
- Opening: strong hook connecting Devon's specific background to THIS role (not generic)
- Middle: 2 concrete examples from Devon's experience that directly match job requirements
  (use specifics: team sizes, technologies, outcomes — not vague claims)
- Closing: clear call to action
- Under 350 words

STYLE RULES:
- Write naturally — this should sound like a real person, not an AI
- NO phrases like: "I am excited to apply", "I am writing to express my interest",
  "I believe I would be a great fit", "I look forward to the opportunity"
- Instead use direct, specific openers like: "Your posting caught my eye because..."
  or "I've spent the last decade solving exactly the kind of problems your team faces."
- Mirror language from the job posting
- Devon's tone: direct, technical, confident but not arrogant

CRITICAL: Start directly with the letter body (Dear Hiring Manager or similar).
No preamble. No "Here is your cover letter." No sign-off commentary after the letter."""

        response = ollama.chat(
            model=MODEL_FAST,
            messages=[{"role":"user","content":prompt}],
            stream=False
        )
        cover_text = response["message"]["content"]

        # Strip LLM preamble/postamble
        cover_text = _strip_llm_noise(cover_text)

        sse_send(task_id, "progress", {"step": "Saving cover letter...", "pct": 85})

        # Save .docx
        from docx import Document as DocxDoc
        from docx.shared import Pt
        safe = re.sub(r"[^\w]","_", company or "Company")
        date_str = datetime.now().strftime("%Y%m%d_%H%M")
        filename = f"CoverLetter_{safe}_{date_str}.docx"
        doc = DocxDoc()
        for line in cover_text.split("\n"):
            p = doc.add_paragraph(line.strip()) if line.strip() else doc.add_paragraph()
            for run in p.runs: run.font.size = Pt(11)
        doc.save(str(COVER_DIR / filename))

        con = get_db()
        con.execute("INSERT INTO cover_letters (name,company,role,filename,created) VALUES (?,?,?,?,?)",
                   (f"Cover Letter — {company}", company, role, filename,
                    datetime.now().strftime("%Y-%m-%d")))
        con.commit(); con.close()

        sse_send(task_id, "done", {
            "content": cover_text,
            "filename": filename,
            "message": f"Cover letter written for {company}"
        })
    except Exception as e:
        sse_send(task_id, "error", {"message": str(e)})


# ── Agent: Scout / Digest ─────────────────────────────────────────────────────

@app.route("/api/agent/scout", methods=["POST"])
def agent_scout():
    task_id = f"scout_{int(time.time())}"
    run_agent_thread(task_id, _run_scout)
    return jsonify({"task_id": task_id})


@app.route("/api/dedup", methods=["POST"])
def api_dedup():
    try:
        from dedup import run_dedup
        result = run_dedup(apply=True, verbose=False)
        return jsonify({
            "message": f"Removed {result['removed']} duplicate jobs across {result['groups']} groups",
            "removed": result["removed"],
            "groups": result["groups"],
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/agent/describe", methods=["POST"])
def agent_describe():
    limit = request.json.get("limit", 30) if request.json else 30
    task_id = f"describe_{int(time.time())}"
    run_agent_thread(task_id, lambda tid: _run_describe(tid, limit))
    return jsonify({"task_id": task_id})


def _run_describe(task_id: str, limit: int = 30):
    sse_send(task_id, "progress", {"step": "Fetching job descriptions...", "pct": 10})
    try:
        from describe import fetch_descriptions_for_jobs
        count = fetch_descriptions_for_jobs(limit=limit, verbose=False)
        sse_send(task_id, "done", {
            "message": f"Fetched {count} job descriptions — re-run Scout to improve scores"
        })
    except Exception as e:
        sse_send(task_id, "error", {"message": str(e)})


def _run_scout(task_id: str):
    sse_send(task_id, "progress", {"step": "Loading new jobs from database...", "pct": 10})
    try:
        con = get_db()
        unscored = con.execute(
            "SELECT COUNT(*) FROM jobs WHERE status='new' AND score IS NULL"
        ).fetchone()[0]
        con.close()

        if not unscored:
            sse_send(task_id, "done", {"message": "No new jobs to score. Run scraper first.", "digest": []})
            return

        sse_send(task_id, "progress", {"step": f"Scoring {min(unscored, 50)} jobs against your profile...", "pct": 30})

        from scorer import score_new_jobs
        digest = score_new_jobs(limit=50, verbose=False)

        sse_send(task_id, "progress", {"step": "Processing scores...", "pct": 90})

        strong = len([d for d in digest if d["score"] >= 8])
        sse_send(task_id, "done", {
            "digest": digest,
            "message": f"Scored {len(digest)} jobs — {strong} strong matches"
        })
    except Exception as e:
        sse_send(task_id, "error", {"message": str(e)})


# ── Agent: Interview Prep ─────────────────────────────────────────────────────

@app.route("/api/agent/interview", methods=["POST"])
def agent_interview():
    d = request.json or {}
    task_id = f"interview_{int(time.time())}"
    run_agent_thread(task_id, _run_interview,
                     d.get("company",""), d.get("role",""), d.get("url",""), d.get("job_id"))
    return jsonify({"task_id": task_id})


def _run_interview(task_id: str, company: str, role: str, url: str, job_id):
    sse_send(task_id, "progress", {"step": "Researching company and role...", "pct": 15})
    try:
        job_text = ""
        if url:
            import requests as req
            from bs4 import BeautifulSoup
            r = req.get(url, headers={"User-Agent":"Mozilla/5.0"}, timeout=15)
            soup = BeautifulSoup(r.text,"html.parser")
            for t in soup(["script","style","nav","footer"]): t.decompose()
            job_text = soup.get_text(separator="\n",strip=True)[:2000]

        sse_send(task_id, "progress", {"step": "Generating interview prep guide...", "pct": 40})

        import ollama
        prompt = f"""Create a comprehensive interview prep guide for Devon O'Rourke interviewing at:
Company: {company}
Role: {role}
{f"Job details: {job_text[:1500]}" if job_text else ""}

Devon's background: E2E architecture, infrastructure, telecom, product roadmap delivery,
lab testing, UAT, staging, pilot programs, executive reporting, professional debugging.

Generate:
## COMPANY RESEARCH
- 3 key facts about {company}'s engineering culture and programs
- Current strategic priorities relevant to TPM work

## LIKELY INTERVIEW QUESTIONS
List 8 behavioral and technical questions likely for this role, with Devon's ideal answer approach for each.
Format: Q: [question] / A: [Devon's angle using his specific background]

## STAR STORIES TO PREPARE
3 specific STAR (Situation/Task/Action/Result) stories Devon should prepare, mapped to this role's requirements.

## KEY TALKING POINTS
5 bullet points Devon should naturally weave into every answer.

## QUESTIONS TO ASK THEM
4 sharp questions Devon should ask the interviewer that show strategic thinking.

## RED FLAGS TO WATCH FOR
2-3 things to listen for that might indicate role/culture mismatch.

Be specific to Devon's actual background. Do not be generic."""

        response = ollama.chat(
            model=MODEL_FAST,
            messages=[{"role":"user","content":prompt}],
            stream=False
        )
        content = response["message"]["content"]

        sse_send(task_id, "progress", {"step": "Saving prep guide...", "pct": 90})

        con = get_db()
        con.execute("INSERT INTO interview_prep (company,role,content,created,job_id) VALUES (?,?,?,?,?)",
                   (company, role, content, datetime.now().strftime("%Y-%m-%d"), job_id))
        con.commit(); con.close()

        sse_send(task_id, "done", {
            "content": content,
            "message": f"Interview prep ready for {company}"
        })
    except Exception as e:
        sse_send(task_id, "error", {"message": str(e)})


# ── Ollama status check ───────────────────────────────────────────────────────

@app.route("/api/ollama/status")
def ollama_status():
    """
    Check Ollama by hitting its REST API directly via requests.
    Avoids CORS issues — browser calls /api/ollama/status on Flask,
    Flask calls localhost:11434, returns result to browser.
    """
    try:
        import requests as req
        r = req.get("http://localhost:11434/api/tags", timeout=3)
        data = r.json()
        names = [m.get("name","") for m in data.get("models", [])]
        return jsonify({"running": True, "models": names})
    except Exception as e:
        # Fallback: try ollama python library
        try:
            import ollama
            models = ollama.list()
            names = [m["name"] for m in models.get("models", [])]
            return jsonify({"running": True, "models": names})
        except:
            return jsonify({"running": False, "models": [], "error": str(e)})


@app.route("/api/ollama/chat", methods=["POST"])
def ollama_chat_proxy():
    """
    Proxy chat requests to Ollama — fixes browser CORS restriction.
    Dashboard sends to /api/ollama/chat, Flask forwards to localhost:11434.
    Supports streaming via SSE.
    """
    import requests as req
    d = request.json or {}
    model   = d.get("model", "llama3.1:8b")
    messages = d.get("messages", [])
    stream  = d.get("stream", False)

    payload = {"model": model, "messages": messages, "stream": stream}

    if stream:
        def generate():
            try:
                r = req.post("http://localhost:11434/api/chat",
                             json=payload, stream=True, timeout=120)
                for line in r.iter_lines():
                    if line:
                        yield line.decode("utf-8") + "\n"
            except Exception as e:
                yield json.dumps({"error": str(e)}) + "\n"
        return Response(generate(), mimetype="application/x-ndjson")
    else:
        try:
            r = req.post("http://localhost:11434/api/chat",
                         json=payload, timeout=120)
            return jsonify(r.json())
        except Exception as e:
            return jsonify({"error": str(e)}), 500


# ── File serving ──────────────────────────────────────────────────────────────

@app.route("/api/resumes/list-files")
def list_resume_files():
    files = [f.name for f in RESUME_DIR.glob("*.docx")]
    return jsonify(files)


# ── Chat endpoint (streams from Ollama) ──────────────────────────────────────

@app.route("/api/chat/test", methods=["GET"])
def chat_test():
    """Debug endpoint — check if Ollama is reachable and what models are available."""
    try:
        import requests as req
        r = req.get("http://localhost:11434/api/tags", timeout=3)
        models = [m.get("name","") for m in r.json().get("models",[])]
        return jsonify({"ok": True, "models": models, "ollama_reachable": True})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e), "ollama_reachable": False})


@app.route("/api/chat", methods=["POST"])
def chat():
    """Stream chat response from Ollama with Devon's profile as system prompt."""
    d = request.json or {}
    messages  = d.get("messages", [])
    model     = d.get("model", "llama3.1:8b")
    job_ctx   = d.get("job_context", "")

    # System prompt — try profile.py first, fall back to built-in
    system = (
        "You are an expert career assistant and coding partner for Devon O'Rourke, "
        "a Senior Technical Program Manager based in Boston, MA.\n\n"
        "Devon's background:\n"
        "- E2E architecture design and delivery\n"
        "- Infrastructure planning and telecom systems\n"
        "- Product roadmap development and cross-functional delivery\n"
        "- Lab testing, UAT, staging, and pilot product rollouts\n"
        "- Executive reporting and board-level presentations\n"
        "- Professional debugging across hardware and software stacks\n"
        "- Target roles: Technical Program Manager, Product Manager, BizOps\n"
        "- Location preference: Remote or Boston/MA area\n\n"
        "Be direct, specific, and practical. When asked about jobs or resumes, "
        "use Devon's actual background above. For coding, write clean working code."
    )
    try:
        from profile import CHAT_SYSTEM_PROMPT
        system = CHAT_SYSTEM_PROMPT
    except Exception:
        pass  # use built-in system prompt above

    if job_ctx:
        system += f"\n\nJob currently open in the dashboard:\n{job_ctx[:1500]}"

    def generate():
        try:
            # Use requests to call Ollama directly — more reliable than the ollama library
            import requests as req
            payload = {
                "model": model,
                "messages": [{"role":"system","content":system}] + messages,
                "stream": True,
                "options": {"num_predict": 1024},
            }
            r = req.post(
                "http://localhost:11434/api/chat",
                json=payload,
                stream=True,
                timeout=120,
            )
            if r.status_code != 200:
                err_msg = f'Ollama returned {r.status_code}: {r.text[:200]}'
                yield 'data: ' + json.dumps({'error': err_msg}) + '\n\n'
                return

            for line in r.iter_lines():
                if not line:
                    continue
                try:
                    chunk = json.loads(line.decode("utf-8"))
                    token = chunk.get("message",{}).get("content","")
                    if token:
                        yield 'data: ' + json.dumps({'token': token}) + '\n\n'
                    if chunk.get("done"):
                        yield 'data: ' + json.dumps({'done': True}) + '\n\n'
                        break
                except Exception:
                    continue

        except Exception as e:
            yield 'data: ' + json.dumps({'error': str(e)}) + '\n\n'

    return Response(
        generate(),
        mimetype="text/event-stream",
        headers={"Cache-Control":"no-cache", "X-Accel-Buffering":"no",
                 "Access-Control-Allow-Origin":"*"},
    )


# ── Admin: read file ──────────────────────────────────────────────────────────

@app.route("/api/admin/files", methods=["GET"])
def admin_list_files():
    """List all editable files with their sizes and modification times."""
    import os
    files = []
    for fname in EDITABLE_FILES:
        path = BASE_DIR / fname
        if path.exists():
            stat = path.stat()
            files.append({
                "name": fname,
                "size": stat.st_size,
                "size_kb": round(stat.st_size / 1024, 1),
                "modified": datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M"),
                "lines": len(path.read_text(encoding="utf-8", errors="ignore").splitlines()),
            })
        else:
            files.append({"name": fname, "size": 0, "size_kb": 0, "modified": "—", "lines": 0, "missing": True})
    return jsonify(files)


@app.route("/api/admin/file", methods=["GET"])
def admin_read_file():
    filename = request.args.get("file","")
    # Uses module-level EDITABLE_FILES
    if filename not in EDITABLE_FILES:
        return jsonify({"error": "File not in allowed list"}), 403
    path = BASE_DIR / filename
    if not path.exists():
        return jsonify({"error": "File not found"}), 404
    return jsonify({"filename": filename, "content": path.read_text(encoding="utf-8")})


@app.route("/api/admin/file", methods=["POST"])
def admin_write_file():
    d = request.json or {}
    filename = d.get("file","")
    content  = d.get("content","")
    # Uses module-level EDITABLE_FILES
    if filename not in EDITABLE_FILES:
        return jsonify({"error": "File not in allowed list"}), 403
    path = BASE_DIR / filename
    # Backup first
    import shutil
    backup_dir = BASE_DIR / "backups"
    backup_dir.mkdir(exist_ok=True)
    if path.exists():
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        shutil.copy(path, backup_dir / f"{filename}.{ts}.bak")
    path.write_text(content, encoding="utf-8")
    return jsonify({"ok": True, "message": f"{filename} saved (backup created)"})


# ── Admin: LLM code edit ──────────────────────────────────────────────────────

@app.route("/api/admin/llm-edit", methods=["POST"])
def admin_llm_edit():
    """Ask LLM to edit a file based on a plain-English instruction."""
    d = request.json or {}
    filename    = d.get("file","scraper.py")
    instruction = d.get("instruction","")
    edit_model  = d.get("model", MODEL_CODE)  # default to deepseek for code edits
    task_id     = f"edit_{int(time.time())}"

    # Uses module-level EDITABLE_FILES
    if filename not in EDITABLE_FILES:
        return jsonify({"error":"File not permitted"}), 403

    path = BASE_DIR / filename
    if not path.exists():
        return jsonify({"error":"File not found"}), 404

    current_code = path.read_text(encoding="utf-8")

    def do_edit(task_id):
        sse_send(task_id, "progress", {"step": f"Reading {filename} ({len(current_code)} chars)...", "pct": 10})
        try:
            import requests as req, re
            prompt = f"""You are an expert developer. Edit this file based on the instruction.

File: {filename}
Instruction: {instruction}

Current file (first 8000 chars):
{current_code[:8000]}

Rules:
- Return ONLY the complete updated file, no explanations, no markdown fences
- Make minimal changes to fulfill the instruction
- Preserve all existing functionality

Updated file:"""

            sse_send(task_id, "progress", {"step": f"Streaming edit from {edit_model}...", "pct": 30})

            # Stream via Ollama REST API directly for real-time progress
            new_code_parts = []
            char_count = 0
            r = req.post(
                "http://localhost:11434/api/chat",
                json={
                    "model": edit_model,
                    "messages": [{"role":"user","content":prompt}],
                    "stream": True,
                    "options": {"temperature": 0.1, "num_predict": 4096},
                },
                stream=True,
                timeout=180,
            )

            for line in r.iter_lines():
                if not line: continue
                try:
                    chunk = json.loads(line.decode("utf-8"))
                    token = chunk.get("message",{}).get("content","")
                    if token:
                        new_code_parts.append(token)
                        char_count += len(token)
                        # Send progress every ~500 chars
                        if char_count % 500 < len(token):
                            pct = min(90, 30 + int((char_count / max(len(current_code), 1000)) * 60))
                            sse_send(task_id, "progress", {
                                "step": f"Generating... ({char_count} chars written)",
                                "pct": pct
                            })
                    if chunk.get("done"):
                        break
                except Exception:
                    continue

            new_code = "".join(new_code_parts).strip()
            # Strip markdown fences
            new_code = re.sub(r"^```[a-z]*\n?", "", new_code, flags=re.MULTILINE)
            if not new_code:
                sse_send(task_id, "error", {"message": "LLM returned empty response — try a shorter file or simpler instruction"})
                return

            sse_send(task_id, "done", {
                "original": current_code,
                "proposed": new_code,
                "filename": filename,
                "instruction": instruction,
                "message": f"Edit ready ({len(new_code)} chars) — review diff before applying",
            })
        except Exception as e:
            sse_send(task_id, "error", {"message": str(e)})

    run_agent_thread(task_id, do_edit)
    return jsonify({"task_id": task_id})


# ── Market Trend Agent ───────────────────────────────────────────────────────

@app.route("/api/agent/market", methods=["POST"])
def agent_market():
    d = request.json or {}
    days = d.get("days", 7)
    task_id = f"market_{int(time.time())}"
    run_agent_thread(task_id, _run_market, days)
    return jsonify({"task_id": task_id})


def _run_market(task_id: str, days: int):
    sse_send(task_id, "progress", {"step": "Analyzing job market data...", "pct": 10})
    try:
        sys.path.insert(0, str(BASE_DIR))
        import importlib
        import market_agent as ma
        importlib.reload(ma)
        sse_send(task_id, "progress", {"step": "Computing trends...", "pct": 40})
        report = ma.run_market_analysis(days_back=days, save=True, silent=True)
        sse_send(task_id, "progress", {"step": "Generating LLM brief...", "pct": 80})
        sse_send(task_id, "done", {
            "message": "Market analysis complete",
            "report": report,
        })
    except Exception as e:
        sse_send(task_id, "error", {"message": str(e)})


@app.route("/api/agent/market/latest", methods=["GET"])
def market_latest():
    """Return latest market report directly."""
    try:
        sys.path.insert(0, str(BASE_DIR))
        import importlib
        import market_agent as ma
        importlib.reload(ma)
        report = ma.run_market_analysis(days_back=7, save=False, silent=True)
        return jsonify(report)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ── Custom Company Targets ───────────────────────────────────────────────────

@app.route("/api/custom-targets", methods=["GET"])
def get_custom_targets():
    try:
        sys.path.insert(0, str(BASE_DIR))
        import importlib
        import custom_targets as ct
        importlib.reload(ct)
        return jsonify(ct.load_targets())
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/custom-targets/add", methods=["POST"])
def add_custom_target():
    d = request.json or {}
    names = [n.strip() for n in d.get("companies","").split(",") if n.strip()]
    try:
        sys.path.insert(0, str(BASE_DIR))
        import importlib
        import custom_targets as ct
        importlib.reload(ct)
        ct.add_companies(names)
        return jsonify({"ok": True, "added": names})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/custom-targets/remove", methods=["POST"])
def remove_custom_target():
    d = request.json or {}
    name = d.get("company","")
    try:
        sys.path.insert(0, str(BASE_DIR))
        import importlib
        import custom_targets as ct
        importlib.reload(ct)
        ct.remove_company(name)
        return jsonify({"ok": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/agent/custom-scrape", methods=["POST"])
def agent_custom_scrape():
    d = request.json or {}
    company = d.get("company","")
    task_id = f"custom_{int(time.time())}"
    run_agent_thread(task_id, _run_custom_scrape, company)
    return jsonify({"task_id": task_id})


def _run_custom_scrape(task_id: str, company: str):
    sse_send(task_id, "progress", {"step": f"Scraping custom targets...", "pct": 10})
    try:
        sys.path.insert(0, str(BASE_DIR))
        import importlib
        import custom_targets as ct
        importlib.reload(ct)
        total = ct.scrape_all_custom(company_filter=company)
        sse_send(task_id, "done", {"message": f"Found {total} new jobs from custom targets"})
    except Exception as e:
        sse_send(task_id, "error", {"message": str(e)})


# ── LinkedIn Scraper (standalone — also runs as part of main scraper) ─────────

@app.route("/api/agent/linkedin", methods=["POST"])
def agent_linkedin():
    """Run LinkedIn scraper standalone. Also runs automatically as part of /api/agent/scrape."""
    d = request.json or {}
    limit = d.get("limit", 30)
    task_id = f"linkedin_{int(time.time())}"
    run_agent_thread(task_id, _run_linkedin, limit)
    return jsonify({"task_id": task_id})


def _run_linkedin(task_id: str, limit: int):
    sse_send(task_id, "progress", {"step": "Starting LinkedIn scraper (browser will open)...", "pct": 10})
    try:
        sys.path.insert(0, str(BASE_DIR))
        import importlib

        # Prefer the integrated scraper.py function (uses same save_jobs, dedup, etc.)
        try:
            import scraper as sc
            importlib.reload(sc)
            sse_send(task_id, "progress", {"step": "Logging into LinkedIn...", "pct": 25})
            jobs = sc.scrape_linkedin()
            if jobs:
                new, updated = sc.save_jobs(jobs)
                sse_send(task_id, "done", {"message": f"LinkedIn: {new} new, {updated} updated jobs"})
            else:
                sse_send(task_id, "done", {"message": "LinkedIn: 0 jobs found (check credentials or session)"})
        except AttributeError:
            # Fallback: use standalone linkedin_scraper.py if scraper.py doesn't have scrape_linkedin yet
            import linkedin_scraper as ls
            importlib.reload(ls)
            sse_send(task_id, "progress", {"step": "Logging into LinkedIn (standalone mode)...", "pct": 25})
            new = ls.run_linkedin_scrape(limit=limit)
            sse_send(task_id, "done", {"message": f"LinkedIn: {new} new jobs added"})
    except Exception as e:
        sse_send(task_id, "error", {"message": str(e)})


@app.route("/api/agent/liveness", methods=["POST"])
def agent_liveness():
    """Check stale aggregator jobs and mark expired ones."""
    d = request.json or {}
    days_old = d.get("days_old", 30)
    limit    = d.get("limit", 300)
    task_id  = f"liveness_{int(time.time())}"
    run_agent_thread(task_id, _run_liveness, days_old, limit)
    return jsonify({"task_id": task_id})


def _run_liveness(task_id: str, days_old: int, limit: int):
    sse_send(task_id, "progress", {"step": f"Checking aggregator jobs older than {days_old} days...", "pct": 5})
    try:
        sys.path.insert(0, str(BASE_DIR))
        import importlib
        import scraper as sc
        importlib.reload(sc)

        def progress(msg):
            pct = 50 if "Done" not in msg else 95
            sse_send(task_id, "progress", {"step": msg.strip(), "pct": pct})

        stats = sc.check_stale_jobs(days_old=days_old, limit=limit, progress_fn=progress)
        sse_send(task_id, "done", {
            "message": (
                f"Liveness check complete: {stats['expired']} expired, "
                f"{stats['live']} live, {stats['uncertain']} uncertain "
                f"(of {stats['checked']} checked)"
            )
        })
    except Exception as e:
        sse_send(task_id, "error", {"message": str(e)})


if __name__ == "__main__":
    init_db()
    print("\n" + "="*55)
    print("  TPM Command Center")
    print("  http://localhost:5050")
    print("="*55 + "\n")
    app.run(host="0.0.0.0", port=5050, debug=False, threaded=True)
