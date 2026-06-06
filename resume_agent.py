"""
TPM Resume Tailor Agent — powered by CrewAI + Ollama
=====================================================
Reads your master resume (.docx), fetches a job posting URL,
then rewrites your resume bullets to match the role.
Saves a tailored .docx and optionally marks the job applied in jobs.db.

Usage:
    python resume_agent.py --job "https://careers.google.com/jobs/..." --resume "resume.docx"
    python resume_agent.py  # interactive mode — prompts for URL and resume path

Install:
    pip install crewai crewai-tools requests beautifulsoup4 python-docx ollama

Author profile baked in:
    Background: E2E architecture, infrastructure, telecom, product roadmap & delivery,
                lab testing, UAT, staging, pilot product testing, executive reporting,
                professional debugging
"""

import argparse
import os
import re
import sqlite3
import sys
from datetime import datetime
from pathlib import Path

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

try:
    from crewai import Agent, Crew, Process, Task
    from crewai.tools import tool
    import ollama as ollama_client
except ImportError:
    print("\n[!] Missing dependencies. Run:")
    print("    pip install crewai crewai-tools requests beautifulsoup4 python-docx ollama\n")
    sys.exit(1)

# ── Config ────────────────────────────────────────────────────────────────────

OLLAMA_MODEL   = "llama3.1:70b-instruct-q4_K_M"
DB_PATH        = "jobs.db"
OUTPUT_DIR     = Path("tailored_resumes")
OUTPUT_DIR.mkdir(exist_ok=True)

# Devon's TPM background — injected into every agent prompt
CANDIDATE_PROFILE = """
Candidate background (Devon O'Rourke — Technical Program Manager):
- End-to-end architecture design and delivery
- Infrastructure planning and implementation
- Telecom systems (network architecture, carrier integrations)
- Product roadmap development and cross-functional delivery
- Lab testing, UAT, staging environments, and pilot product rollouts
- Executive reporting, stakeholder communication, and board-level presentations
- Professional debugging — complex system triage across hardware and software stacks
- Strong experience managing ambiguity, driving alignment, and unblocking engineering teams
"""

# ── Tools available to agents ─────────────────────────────────────────────────

@tool("fetch_job_posting")
def fetch_job_posting(url: str) -> str:
    """Fetches and extracts the text content of a job posting from a URL."""
    try:
        headers = {
            "User-Agent": (
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                "AppleWebKit/537.36 (KHTML, like Gecko) "
                "Chrome/122.0.0.0 Safari/537.36"
            )
        }
        r = requests.get(url, headers=headers, timeout=15)
        soup = BeautifulSoup(r.text, "html.parser")

        # Remove nav, footer, scripts
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()

        text = soup.get_text(separator="\n", strip=True)
        # Trim to ~4000 chars — enough for any job description
        return text[:4000]
    except Exception as e:
        return f"Error fetching job posting: {e}"


@tool("read_resume")
def read_resume(path: str) -> str:
    """Reads a .docx resume file and returns its full text content."""
    try:
        doc = Document(path)
        lines = []
        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text.strip())
        return "\n".join(lines)
    except Exception as e:
        return f"Error reading resume: {e}"


@tool("save_tailored_resume")
def save_tailored_resume(content: str, company: str, role: str, original_path: str) -> str:
    """
    Saves the tailored resume content as a new .docx file.
    Content should be formatted with sections separated by '##' headers.
    """
    try:
        # Load original to preserve formatting structure
        original = Document(original_path)
        new_doc  = Document()

        # Set margins (1 inch)
        from docx.oxml.ns import qn
        from docx.oxml   import OxmlElement
        section = new_doc.sections[0]
        section.top_margin    = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin   = Pt(72)
        section.right_margin  = Pt(72)

        # Header — candidate name from original doc first paragraph
        name_text = original.paragraphs[0].text if original.paragraphs else "Devon O'Rourke"
        heading = new_doc.add_paragraph(name_text)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.runs[0] if heading.runs else heading.add_run(name_text)
        run.bold      = True
        run.font.size = Pt(16)

        # Tailored label
        label = new_doc.add_paragraph(f"Tailored for: {role} at {company}")
        label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_run = label.runs[0] if label.runs else label.add_run()
        label_run.font.size  = Pt(10)
        label_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        new_doc.add_paragraph()  # spacer

        # Parse and write the tailored content
        for line in content.split("\n"):
            line = line.strip()
            if not line:
                new_doc.add_paragraph()
                continue

            if line.startswith("## "):
                # Section heading
                p = new_doc.add_paragraph(line[3:].upper())
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(11)
                # Underline via border
                pPr = p._p.get_or_add_pPr()
                pBdr = OxmlElement("w:pBdr")
                bottom = OxmlElement("w:bottom")
                bottom.set(qn("w:val"), "single")
                bottom.set(qn("w:sz"), "6")
                bottom.set(qn("w:space"), "1")
                bottom.set(qn("w:color"), "000000")
                pBdr.append(bottom)
                pPr.append(pBdr)

            elif line.startswith("- ") or line.startswith("• "):
                # Bullet point
                p = new_doc.add_paragraph(style="List Bullet")
                p.add_run(line[2:])
                p.runs[0].font.size = Pt(10)

            elif line.startswith("**") and line.endswith("**"):
                # Bold sub-heading (job title / company)
                p = new_doc.add_paragraph(line.strip("*"))
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(10)

            else:
                p = new_doc.add_paragraph(line)
                p.runs[0].font.size = Pt(10) if p.runs else None

        # Save file
        safe_company = re.sub(r"[^\w\-]", "_", company)
        safe_role    = re.sub(r"[^\w\-]", "_", role)[:30]
        date_str     = datetime.now().strftime("%Y%m%d")
        filename     = OUTPUT_DIR / f"Resume_{safe_company}_{safe_role}_{date_str}.docx"
        new_doc.save(str(filename))
        return f"Saved: {filename}"

    except Exception as e:
        return f"Error saving resume: {e}"


@tool("mark_job_applied")
def mark_job_applied(url: str) -> str:
    """Marks a job posting as 'applied' in the jobs tracker database."""
    try:
        con = sqlite3.connect(DB_PATH)
        cur = con.execute(
            "UPDATE jobs SET status = 'applied' WHERE url = ?", (url,)
        )
        con.commit()
        updated = cur.rowcount
        con.close()
        if updated:
            return f"Marked as applied in tracker: {url}"
        return f"Job URL not found in tracker (not an error — may have been added manually): {url}"
    except Exception as e:
        return f"DB error: {e}"


# ── Agents ────────────────────────────────────────────────────────────────────

def build_crew(job_url: str, resume_path: str) -> Crew:

    llm_config = {
        "model": f"ollama/{OLLAMA_MODEL}",
        "base_url": "http://localhost:11434",
    }

    # Agent 1 — Job Analyst
    job_analyst = Agent(
        role="Senior Technical Recruiter & Job Analyst",
        goal=(
            "Deeply analyze a job posting and extract every requirement, "
            "keyword, skill, and cultural signal that matters for a TPM role."
        ),
        backstory=(
            "You are a former FAANG technical recruiter who has reviewed thousands "
            "of TPM applications. You know exactly what hiring managers look for "
            "and how ATS systems score resumes against job descriptions. "
            "You are ruthlessly precise about keyword matching."
        ),
        tools=[fetch_job_posting],
        llm=llm_config,
        verbose=True,
        allow_delegation=False,
    )

    # Agent 2 — Resume Writer
    resume_writer = Agent(
        role="Expert TPM Resume Writer",
        goal=(
            "Rewrite the candidate's resume bullets to maximally match the target "
            "job posting while staying 100% truthful to their actual experience."
        ),
        backstory=(
            f"You are an elite resume writer specializing in Technical Program Manager "
            f"roles at top tech companies. You transform generic resume bullets into "
            f"powerful, ATS-optimized, achievement-focused statements that get interviews. "
            f"You NEVER fabricate experience — you reframe and reword what's real.\n\n"
            f"{CANDIDATE_PROFILE}"
        ),
        tools=[read_resume, save_tailored_resume],
        llm=llm_config,
        verbose=True,
        allow_delegation=False,
    )

    # Agent 3 — Application Coordinator
    coordinator = Agent(
        role="Application Coordinator",
        goal="Track the job application in the database and provide a final summary.",
        backstory=(
            "You keep the job search organized. Once a resume is tailored and saved, "
            "you update the tracker and give the candidate a clear action summary."
        ),
        tools=[mark_job_applied],
        llm=llm_config,
        verbose=True,
        allow_delegation=False,
    )

    # ── Tasks ──────────────────────────────────────────────────────────────────

    task_analyze = Task(
        description=(
            f"Fetch and analyze the job posting at: {job_url}\n\n"
            "Extract and return:\n"
            "1. Company name and exact role title\n"
            "2. Top 10 required skills and keywords (especially technical ones)\n"
            "3. Top 5 preferred/bonus qualifications\n"
            "4. Key responsibilities (summarized)\n"
            "5. Any specific tools, methodologies, or frameworks mentioned\n"
            "6. Tone and culture signals (startup vs enterprise, speed vs process, etc.)\n"
            "7. ATS keywords to prioritize in the resume"
        ),
        expected_output=(
            "A structured analysis with: company, role title, required skills list, "
            "preferred skills list, key responsibilities, tools/frameworks, culture notes, "
            "and a prioritized ATS keyword list."
        ),
        agent=job_analyst,
    )

    task_rewrite = Task(
        description=(
            f"Read the resume at: {resume_path}\n\n"
            "Using the job analysis from the previous task, rewrite the resume to:\n"
            "1. Mirror the exact language and keywords from the job posting\n"
            "2. Lead each bullet with a strong action verb\n"
            "3. Add quantifiable metrics wherever the candidate's background supports it\n"
            "   (e.g., 'reduced deployment time by X%', 'managed X-engineer teams', etc.)\n"
            "4. Prioritize experience most relevant to this specific role\n"
            "5. Ensure E2E architecture, infrastructure, telecom, lab testing, UAT, "
            "   staging, pilot testing, and executive reporting are highlighted "
            "   where they match the job requirements\n"
            "6. Keep the resume to 1 page if possible, 2 pages maximum\n\n"
            "Format the output with ## section headers and - bullet points.\n"
            "Then save it using the save_tailored_resume tool with the company name and role."
        ),
        expected_output=(
            "A complete tailored resume in formatted text, plus confirmation that "
            "the .docx file was saved successfully with the file path."
        ),
        agent=resume_writer,
        context=[task_analyze],
    )

    task_track = Task(
        description=(
            f"Mark the job at {job_url} as 'applied' in the tracker database.\n"
            "Then provide a final summary including:\n"
            "1. Company and role applied to\n"
            "2. Path to the saved tailored resume\n"
            "3. Top 3 keywords to emphasize if asked about this role in an interview\n"
            "4. One suggested talking point based on Devon's background that strongly "
            "   matches this specific role"
        ),
        expected_output=(
            "Confirmation of tracker update, resume file path, top keywords, "
            "and one strong interview talking point."
        ),
        agent=coordinator,
        context=[task_analyze, task_rewrite],
    )

    return Crew(
        agents=[job_analyst, resume_writer, coordinator],
        tasks=[task_analyze, task_rewrite, task_track],
        process=Process.sequential,
        verbose=True,
    )


# ── Entry point ───────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="TPM Resume Tailor Agent")
    parser.add_argument("--job",    type=str, help="Job posting URL")
    parser.add_argument("--resume", type=str, help="Path to your master resume .docx")
    args = parser.parse_args()

    print("\n" + "═"*55)
    print("  TPM Resume Tailor Agent")
    print("  Powered by CrewAI + Ollama (llama3.1:70b)")
    print("═"*55 + "\n")

    # Get job URL
    job_url = args.job
    if not job_url:
        job_url = input("Paste the job posting URL: ").strip()
    if not job_url:
        print("No URL provided. Exiting.")
        sys.exit(1)

    # Get resume path
    resume_path = args.resume
    if not resume_path:
        resume_path = input("Path to your master resume .docx [resume.docx]: ").strip()
        if not resume_path:
            resume_path = "resume.docx"

    if not Path(resume_path).exists():
        print(f"\n[!] Resume file not found: {resume_path}")
        print("    Place your resume.docx in the same folder as this script,")
        print("    or pass the full path with --resume\n")
        sys.exit(1)

    # Check Ollama is running
    try:
        ollama_client.list()
    except Exception:
        print("\n[!] Ollama is not running. Open the Ollama desktop app first.\n")
        sys.exit(1)

    print(f"  Job URL   : {job_url}")
    print(f"  Resume    : {resume_path}")
    print(f"  Model     : {OLLAMA_MODEL}")
    print(f"  Output dir: {OUTPUT_DIR}/\n")
    print("Starting agents...\n")

    crew   = build_crew(job_url, resume_path)
    result = crew.kickoff()

    print("\n" + "═"*55)
    print("  DONE")
    print("═"*55)
    print(result)
    print(f"\nTailored resumes saved in: {OUTPUT_DIR}/\n")


if __name__ == "__main__":
    main()
